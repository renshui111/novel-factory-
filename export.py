# export.py — 多格式导出
# 支持 .txt, .epub, .pdf, .docx, .html

import os
import re

try:
    from ebooklib import epub
    HAS_EPUB = True
except ImportError:
    HAS_EPUB = False

try:
    from fpdf2 import FPDF
    HAS_FPDF = True
except ImportError:
    try:
        from fpdf import FPDF
        HAS_FPDF = True
    except ImportError:
        HAS_FPDF = False

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

from core import read_file, write_file, ensure_dir, count_words


def _get_chapters(book_dir: str) -> list:
    """读取账单章节列表"""
    chapter_dir = os.path.join(book_dir, "正文")
    if not os.path.exists(chapter_dir):
        return []
    files = sorted([f for f in os.listdir(chapter_dir) if f.endswith('.md')])
    chapters = []
    for fname in files:
        title = fname.replace('.md', '')
        content = read_file(os.path.join(chapter_dir, fname))
        chapters.append({"num": len(chapters) + 1, "title": title, "content": content})
    return chapters


def _book_name(book_dir: str) -> str:
    """从目录名推断书名"""
    return os.path.basename(book_dir)


def export_to_txt(book_dir: str, output_path: str = "", log_callback=None) -> str:
    """拼接每章 .md 为单个 .txt"""
    chapters = _get_chapters(book_dir)
    if not chapters:
        if log_callback:
            log_callback("未找到正文目录，跳过导出")
        return ""

    if not output_path:
        output_path = os.path.join(book_dir, f"{_book_name(book_dir)}.txt")

    lines = []
    for ch in chapters:
        lines.append(f"\n{'='*60}\n{ch['title']}\n{'='*60}\n")
        lines.append(ch['content'])

    full_text = "\n".join(lines)
    write_file(output_path, full_text)

    wc = count_words(full_text)
    if log_callback:
        log_callback(f"导出 TXT: {output_path} ({wc} 字)")

    return output_path


def export_to_epub(book_dir: str, output_path: str = "", log_callback=None) -> str:
    """生成 .epub 文件"""
    if not HAS_EPUB:
        if log_callback:
            log_callback("ebooklib 未安装，跳过 epub 导出")
        return ""

    chapters = _get_chapters(book_dir)
    if not chapters:
        if log_callback:
            log_callback("未找到正文目录")
        return ""

    book_name = _book_name(book_dir)
    if not output_path:
        output_path = os.path.join(book_dir, f"{book_name}.epub")

    book = epub.EpubBook()
    book.set_identifier(book_name)
    book.set_title(book_name)
    book.set_language('zh-CN')

    chapters_epub = []
    for ch in chapters:
        html_content = ch['content'].replace('\n', '</p><p>')
        chap = epub.EpubHtml(
            title=ch['title'],
            file_name=f'chap_{ch["num"]:03d}.xhtml',
            lang='zh-CN'
        )
        chap.content = f'<h1>{ch["title"]}</h1><p>{html_content}</p>'
        book.add_item(chap)
        chapters_epub.append(chap)

    book.toc = chapters_epub
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())

    style = 'BODY { font-family: "Microsoft YaHei", serif; font-size: 1.1em; }'
    css = epub.EpubItem(
        uid="style",
        file_name="style/default.css",
        media_type="text/css",
        content=style.encode('utf-8'))
    book.add_item(css)
    for ch in chapters_epub:
        ch.add_item(css)

    book.spine = ['nav'] + chapters_epub
    epub.write_epub(output_path, book, {})

    if log_callback:
        log_callback(f"导出 EPUB: {output_path} ({len(chapters)} 章)")

    return output_path


def _md_to_html_simple(text: str) -> str:
    """将简单 Markdown 转为 HTML"""
    lines = text.split('\n')
    html_parts = []
    in_para = False
    for line in lines:
        stripped = line.strip()
        if not stripped:
            if in_para:
                html_parts.append('</p>')
                in_para = False
            continue
        # Headers
        if stripped.startswith('### '):
            if in_para:
                html_parts.append('</p>')
                in_para = False
            html_parts.append(f'<h3>{stripped[4:]}</h3>')
        elif stripped.startswith('## '):
            if in_para:
                html_parts.append('</p>')
                in_para = False
            html_parts.append(f'<h2>{stripped[3:]}</h2>')
        elif stripped.startswith('# '):
            if in_para:
                html_parts.append('</p>')
                in_para = False
            html_parts.append(f'<h1>{stripped[2:]}</h1>')
        else:
            if not in_para:
                html_parts.append('<p>')
                in_para = True
            else:
                html_parts.append('<br/>')
            # Bold/italic
            line_html = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', stripped)
            line_html = re.sub(r'\*(.*?)\*', r'<em>\1</em>', line_html)
            html_parts.append(line_html)
    if in_para:
        html_parts.append('</p>')
    return '\n'.join(html_parts)


def export_to_html(book_dir: str, output_path: str = "", log_callback=None) -> str:
    """生成单文件 HTML，含样式，可浏览器打开阅读"""
    chapters = _get_chapters(book_dir)
    if not chapters:
        if log_callback:
            log_callback("未找到正文目录，跳过导出")
        return ""

    book_name = _book_name(book_dir)
    if not output_path:
        output_path = os.path.join(book_dir, f"{book_name}.html")

    # 封面页
    body_parts = []
    body_parts.append(
        f'<div class="cover">'
        f'<h1 class="book-title">{book_name}</h1>'
        f'<p class="book-meta">{len(chapters)} 章</p>'
        f'</div>'
        f'<hr/>'
    )

    for ch in chapters:
        body_parts.append(f'<div class="chapter">')
        body_parts.append(f'<h2 class="chapter-title">{ch["title"]}</h2>')
        body_parts.append(f'<div class="content">{_md_to_html_simple(ch["content"])}</div>')
        body_parts.append(f'</div>')

    html = f'''<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{book_name}</title>
<style>
* {{ margin: 0; padding: 0; box-sizing: border-box; }}
body {{
    font-family: "Microsoft YaHei", "Noto Sans SC", "Source Han Sans", sans-serif;
    background: #f5f0e8;
    color: #2c2c2c;
    line-height: 1.8;
    max-width: 800px;
    margin: 0 auto;
    padding: 20px;
}}
.cover {{
    text-align: center;
    padding: 80px 20px 40px;
}}
.book-title {{
    font-size: 2.4em;
    color: #8b4513;
    margin-bottom: 10px;
}}
.book-meta {{
    font-size: 1.1em;
    color: #888;
    margin-top: 20px;
}}
hr {{ border: none; border-top: 1px solid #ddd; margin: 30px 0; }}
.chapter {{
    margin-bottom: 40px;
    padding-bottom: 20px;
}}
.chapter-title {{
    font-size: 1.8em;
    color: #333;
    text-align: center;
    margin: 30px 0 20px;
    padding-bottom: 10px;
    border-bottom: 2px solid #e0d5c7;
}}
.content p {{ margin: 0.8em 0; text-indent: 2em; }}
.content h1, .content h2, .content h3 {{
    margin: 1.2em 0 0.6em;
    color: #444;
}}
.content h1 {{ font-size: 1.6em; }}
.content h2 {{ font-size: 1.4em; }}
.content h3 {{ font-size: 1.2em; }}
@media (prefers-color-scheme: dark) {{
    body {{ background: #1a1a1a; color: #d4d4d4; }}
    .book-title {{ color: #e0a050; }}
    .chapter-title {{ color: #ddd; border-bottom-color: #444; }}
    hr {{ border-top-color: #333; }}
    .content h1, .content h2, .content h3 {{ color: #ccc; }}
}}
</style>
</head>
<body>
{chr(10).join(body_parts)}
</body>
</html>'''

    write_file(output_path, html)
    if log_callback:
        wc = count_words("".join(ch["content"] for ch in chapters))
        log_callback(f"导出 HTML: {output_path} ({wc} 字, {len(chapters)} 章)")

    return output_path


def _find_cjk_font() -> str:
    """寻找系统中可用的 CJK Unicode TTF 字体路径"""
    import platform
    if platform.system() == "Windows":
        windir = os.environ.get("WINDIR", "C:\\Windows")
        candidates = [
            os.path.join(windir, "Fonts", f)
            for f in [
                "msyh.ttc",
                "msyhbd.ttc",
                "simsun.ttc",
                "simhei.ttf",
                "yahei.ttf",
                "msyhl.ttc",
            ]
        ]
        for path in candidates:
            if os.path.exists(path):
                return path
    # MacOS
    mac_candidates = [
        "/System/Library/Fonts/PingFang.ttc",
        "/System/Library/Fonts/STHeiti Light.ttc",
        "/Library/Fonts/Arial Unicode.ttf",
    ]
    for path in mac_candidates:
        if os.path.exists(path):
            return path
    # Linux
    import subprocess
    try:
        result = subprocess.run(
            ["fc-match", "-f", "%{file}", "sans-serif"],
            capture_output=True, text=True, timeout=5)
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout.strip()
    except Exception:
        pass
    return ""


def export_to_pdf(book_dir: str, output_path: str = "", log_callback=None) -> str:
    """用 fpdf2 将小说转为 PDF，拼接每章内容，生成带封面的 PDF"""
    if not HAS_FPDF:
        if log_callback:
            log_callback("fpdf2 未安装，跳过 PDF 导出（请执行: pip install fpdf2）")
        return ""

    chapters = _get_chapters(book_dir)
    if not chapters:
        if log_callback:
            log_callback("未找到正文目录，跳过导出")
        return ""

    book_name = _book_name(book_dir)
    if not output_path:
        output_path = os.path.join(book_dir, f"{book_name}.pdf")

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)

    # 添加 Unicode 中文字体
    cjk_font_path = _find_cjk_font()
    has_cjk = bool(cjk_font_path)
    if has_cjk:
        pdf.add_font("CJK", "", cjk_font_path, uni=True)
        # 尝试加粗版本
        bold_path = cjk_font_path.replace("msyh.ttc", "msyhbd.ttc")
        if os.path.exists(bold_path):
            pdf.add_font("CJK", "B", bold_path, uni=True)
        else:
            pdf.add_font("CJK", "B", cjk_font_path, uni=True)

    def set_cjk_font(style="", size=11):
        if has_cjk:
            pdf.set_font("CJK", style, size)
        else:
            pdf.set_font("Helvetica", style, size)

    # --- 封面 ---
    pdf.add_page()
    pdf.ln(60)
    set_cjk_font("B", 28)
    pdf.set_text_color(60, 60, 60)
    pdf.cell(0, 15, book_name, new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(10)
    set_cjk_font("", 14)
    pdf.set_text_color(120, 120, 120)
    pdf.cell(0, 10, f"{len(chapters)} 章", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(20)
    pdf.set_draw_color(180, 180, 180)
    pdf.line(60, pdf.get_y(), 150, pdf.get_y())
    pdf.ln(10)
    set_cjk_font("", 10)
    pdf.set_text_color(160, 160, 160)
    pdf.cell(0, 10, "Generated by Novel Factory", new_x="LMARGIN", new_y="NEXT", align="C")

    # --- 正文 ---
    for ch in chapters:
        pdf.add_page()
        # 章节标题
        set_cjk_font("B", 16)
        pdf.set_text_color(40, 40, 40)
        pdf.cell(0, 12, ch['title'], new_x="LMARGIN", new_y="NEXT", align="C")
        pdf.ln(5)
        pdf.set_draw_color(200, 200, 200)
        pdf.line(30, pdf.get_y(), 180, pdf.get_y())
        pdf.ln(8)

        # 内容
        set_cjk_font("", 11)
        pdf.set_text_color(50, 50, 50)

        # 简单将 Markdown 段落转为 PDF
        paragraphs = ch['content'].split('\n')
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            # 处理标题
            if para.startswith('### '):
                set_cjk_font("B", 12)
                pdf.multi_cell(0, 7, para[4:])
                set_cjk_font("", 11)
                pdf.ln(2)
            elif para.startswith('## '):
                set_cjk_font("B", 13)
                pdf.multi_cell(0, 8, para[3:])
                set_cjk_font("", 11)
                pdf.ln(2)
            elif para.startswith('# '):
                set_cjk_font("B", 14)
                pdf.multi_cell(0, 9, para[2:])
                set_cjk_font("", 11)
                pdf.ln(2)
            else:
                # 移除粗体/斜体标记
                clean = re.sub(r'\*\*(.*?)\*\*', r'\1', para)
                clean = re.sub(r'\*(.*?)\*', r'\1', clean)
                pdf.multi_cell(0, 6, clean)
                pdf.ln(1)

    pdf.output(output_path)

    wc = count_words("".join(ch["content"] for ch in chapters))
    if log_callback:
        log_callback(f"导出 PDF: {output_path} ({wc} 字, {len(chapters)} 章)")

    return output_path


def export_to_docx(book_dir: str, output_path: str = "", log_callback=None) -> str:
    """用 python-docx 生成 .docx 格式，包含标题和章节"""
    if not HAS_DOCX:
        if log_callback:
            log_callback("python-docx 未安装，跳过 DOCX 导出（请执行: pip install python-docx）")
        return ""

    chapters = _get_chapters(book_dir)
    if not chapters:
        if log_callback:
            log_callback("未找到正文目录，跳过导出")
        return ""

    book_name = _book_name(book_dir)
    if not output_path:
        output_path = os.path.join(book_dir, f"{book_name}.docx")

    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # --- 封面标题 ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"\n\n\n{book_name}")
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x8B, 0x45, 0x13)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{len(chapters)} 章")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()

    # --- 正文 ---
    for ch in chapters:
        # 章节标题
        doc.add_heading(ch['title'], level=1)

        # 内容段落
        lines = ch['content'].split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
            if line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            else:
                # 移除粗体/斜体标记（docx 不支持 inline markdown 转换）
                clean = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
                clean = re.sub(r'\*(.*?)\*', r'\1', clean)
                doc.add_paragraph(clean)

    doc.save(output_path)

    wc = count_words("".join(ch["content"] for ch in chapters))
    if log_callback:
        log_callback(f"导出 DOCX: {output_path} ({wc} 字, {len(chapters)} 章)")

    return output_path


def _get_book_info(book_dir: str) -> dict:
    """返回书籍信息：字数、章节数"""
    chapters = _get_chapters(book_dir)
    if not chapters:
        return {"chapters": 0, "words": 0}
    total_words = count_words("".join(ch["content"] for ch in chapters))
    return {"chapters": len(chapters), "words": total_words}


def export_all(book_dir: str, log_callback=None):
    """导出所有支持格式"""
    export_to_txt(book_dir, log_callback=log_callback)
    export_to_epub(book_dir, log_callback=log_callback)
    export_to_pdf(book_dir, log_callback=log_callback)
    export_to_docx(book_dir, log_callback=log_callback)
    export_to_html(book_dir, log_callback=log_callback)